import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders.
    Args:
        cell: docx.table._Cell object
        kwargs: top, bottom, left, right, etc. dicts containing 'sz', 'val', 'color'
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn(f'w:{key}'), str(val))

def apply_text_runs(paragraph, text):
    """
    Parses inline markdown like **bold** and *italic* and adds runs to the paragraph.
    """
    # Regex to extract bold/italic/plain segments
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|[^*`]+)')
    matches = pattern.findall(text)
    
    for part in matches:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def headings_match(h1: str, h2: str) -> bool:
    """Check if two heading texts are conceptually identical to prevent duplicate title rendering."""
    def clean(s):
        # Remove digits, dots, spaces, hyphens, colons, and convert to lowercase
        return re.sub(r'[\d\.\s\-\:]+', '', s.lower()).strip()
    c1, c2 = clean(h1), clean(h2)
    return c1 == c2 or c1 in c2 or c2 in c1

def build_academic_document(title: str, sections: list, output_path: str):
    """
    Generates a word document (.docx) with Minimalist Academic Styling.
    Args:
        title (str): Document title.
        sections (list): List of dicts, e.g. [{"heading": "Introduction", "content": "..."}]
        output_path (str): File path to save the generated document.
    """
    doc = Document()
    
    # Page Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Configuration
    # Normal Style (Body text)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(51, 51, 51) # Charcoal
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.space_before = Pt(0)
    
    # Document Title (Large centered)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_p.paragraph_format.space_after = Pt(24)
    title_p.paragraph_format.space_before = Pt(12)
    
    # Add an empty paragraph or divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(24)
    # A simple clean horizontal line
    p_div_run = p_div.add_run("—" * 30)
    p_div_run.font.color.rgb = RGBColor(150, 150, 150)
    
    for sec in sections:
        heading_text = sec.get("heading", "").strip()
        content_text = sec.get("content", "").strip()
        
        # Add heading
        if heading_text:
            h_p = doc.add_paragraph()
            h_p.paragraph_format.space_before = Pt(18)
            h_p.paragraph_format.space_after = Pt(6)
            h_p.paragraph_format.keep_with_next = True
            
            h_run = h_p.add_run(heading_text)
            h_run.font.name = 'Times New Roman'
            h_run.font.size = Pt(16)
            h_run.font.bold = True
            h_run.font.color.rgb = RGBColor(0, 0, 0)
            
        # Parse content body line by line
        lines = content_text.split('\n')
        in_list = False
        in_table = False
        table_rows = []
        
        for line in lines:
            stripped = line.strip()
            
            # Clean raw markdown link syntax: [link text](url) -> "link text"
            stripped = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', stripped)
            # Remove escape backslashes from markdown symbols (e.g. \* or \- or \_)
            stripped = re.sub(r'\\([*_\-`#])', r'\1', stripped)
            
            # 1. Handle Tables
            if stripped.startswith('|'):
                in_table = True
                # Parse cells
                row_cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
                table_rows.append(row_cells)
                continue
            elif in_table:
                # We just exited a table block or hit a non-table line
                # Let's render the collected table rows
                if len(table_rows) > 0:
                    # Filter out separators (e.g. |---|---|)
                    filtered_rows = []
                    for r in table_rows:
                        if all(re.match(r'^[-:]+$', cell) for cell in r if cell):
                            continue
                        filtered_rows.append(r)
                    
                    if len(filtered_rows) > 0:
                        cols_count = max(len(r) for r in filtered_rows)
                        table = doc.add_table(rows=0, cols=cols_count)
                        table.autofit = True
                        
                        for i, r in enumerate(filtered_rows):
                            row = table.add_row()
                            for j, val in enumerate(r):
                                if j < len(row.cells):
                                    cell = row.cells[j]
                                    cell.text = val
                                    
                                    # Styling the cells (APA/Academic style)
                                    # Light borders for top/bottom of headers, bottom of last row
                                    border_kwargs = {}
                                    if i == 0:
                                        # Header row: top and bottom borders
                                        border_kwargs['top'] = {"sz": 6, "val": "single", "color": "000000"}
                                        border_kwargs['bottom'] = {"sz": 6, "val": "single", "color": "000000"}
                                        # Bold header text
                                        for p in cell.paragraphs:
                                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                            for run in p.runs:
                                                run.font.bold = True
                                                run.font.name = 'Times New Roman'
                                    elif i == len(filtered_rows) - 1:
                                        # Last row: bottom border
                                        border_kwargs['bottom'] = {"sz": 6, "val": "single", "color": "000000"}
                                    
                                    # Clear vertical borders
                                    border_kwargs['left'] = {"sz": 0, "val": "none", "color": "auto"}
                                    border_kwargs['right'] = {"sz": 0, "val": "none", "color": "auto"}
                                    border_kwargs['insideV'] = {"sz": 0, "val": "none", "color": "auto"}
                                    
                                    set_cell_border(cell, **border_kwargs)
                                    
                        # Add space after table
                        doc.add_paragraph().paragraph_format.space_before = Pt(6)
                
                in_table = False
                table_rows = []
                # Don't skip processing the current line if it's not a table line
                if not stripped:
                    continue
            
            # 2. Skip empty lines
            if not stripped:
                continue
                
            # 3. Handle subheadings within section (#, ##, ###, ####)
            if stripped.startswith('# '):
                title_text = stripped[2:].strip()
                if headings_match(heading_text, title_text):
                    continue  # Skip duplication of section heading
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_before = Pt(14)
                sub_p.paragraph_format.space_after = Pt(4)
                sub_p.paragraph_format.keep_with_next = True
                sub_run = sub_p.add_run(title_text)
                sub_run.font.name = 'Times New Roman'
                sub_run.font.size = Pt(14)
                sub_run.font.bold = True
                sub_run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            elif stripped.startswith('## '):
                title_text = stripped[3:].strip()
                if headings_match(heading_text, title_text):
                    continue  # Skip duplication of section heading
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_before = Pt(12)
                sub_p.paragraph_format.space_after = Pt(4)
                sub_p.paragraph_format.keep_with_next = True
                sub_run = sub_p.add_run(title_text)
                sub_run.font.name = 'Times New Roman'
                sub_run.font.size = Pt(13)
                sub_run.font.bold = True
                sub_run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            elif stripped.startswith('### '):
                title_text = stripped[4:].strip()
                if headings_match(heading_text, title_text):
                    continue  # Skip duplication
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_before = Pt(12)
                sub_p.paragraph_format.space_after = Pt(4)
                sub_p.paragraph_format.keep_with_next = True
                sub_run = sub_p.add_run(title_text)
                sub_run.font.name = 'Times New Roman'
                sub_run.font.size = Pt(12)
                sub_run.font.bold = True
                sub_run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            elif stripped.startswith('#### '):
                title_text = stripped[5:].strip()
                if headings_match(heading_text, title_text):
                    continue  # Skip duplication
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_before = Pt(10)
                sub_p.paragraph_format.space_after = Pt(4)
                sub_p.paragraph_format.keep_with_next = True
                sub_run = sub_p.add_run(title_text)
                sub_run.font.name = 'Times New Roman'
                sub_run.font.size = Pt(11)
                sub_run.font.bold = True
                sub_run.font.italic = True
                sub_run.font.color.rgb = RGBColor(0, 0, 0)
                continue
                
            # 4. Handle List Items
            # Bullet list
            bullet_match = re.match(r'^[\-\*]\s+(.*)', stripped)
            # Numbered list
            num_match = re.match(r'^\d+\.\s+(.*)', stripped)
            
            if bullet_match:
                item_text = bullet_match.group(1)
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.5)
                apply_text_runs(p, item_text)
                in_list = True
            elif num_match:
                item_text = num_match.group(1)
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.5)
                apply_text_runs(p, item_text)
                in_list = True
            else:
                # Regular paragraph
                p = doc.add_paragraph()
                apply_text_runs(p, stripped)
                in_list = False
                
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path
